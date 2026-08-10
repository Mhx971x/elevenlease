from pathlib import Path
from datetime import date
from lxml import html
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
DIST = ROOT / "astro-site" / "dist"
OUT = Path(__file__).resolve().parent
DOCX = OUT / "seo-audit-elevenlease-fr-2026-07-24.docx"

NAVY = "1B2A4A"
BLUE = "2563EB"
GREEN = "16A34A"
AMBER = "D97706"
RED = "DC2626"
LIGHT_BLUE = "EFF6FF"
LIGHT_GRAY = "F8F9FA"
BORDER = "E2E8F0"
DARK = "1E293B"
WHITE = "FFFFFF"


def shade(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def borders(table, color=BORDER, size="4"):
    tbl_pr = table._tbl.tblPr
    node = tbl_pr.find(qn("w:tblBorders"))
    if node is None:
        node = OxmlElement("w:tblBorders")
        tbl_pr.append(node)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:color"), color)
        node.append(el)


def margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, *, bold=False, color=DARK, size=9.5, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    margins(cell)


def add_table(doc, headers, rows, widths=None, status_col=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    borders(table)
    for i, head in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], head, bold=True, color=WHITE, size=9)
        shade(table.rows[0].cells[i], NAVY)
        if widths:
            table.rows[0].cells[i].width = Inches(widths[i])
    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=8.8)
            if widths:
                cells[i].width = Inches(widths[i])
            if ri % 2:
                shade(cells[i], LIGHT_GRAY)
        if status_col is not None:
            value = str(row[status_col]).lower()
            color = GREEN if value in ("good", "fort", "strong") else RED if value in ("missing", "critique", "needs work") else AMBER
            shade(cells[status_col], color)
            set_cell_text(cells[status_col], row[status_col], bold=True, color=WHITE, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def page_audit():
    pages = []
    for path in sorted(DIST.rglob("*.html")):
        if path.name == "admin.html":
            continue
        tree = html.fromstring(path.read_text(encoding="utf-8"))
        title = (tree.xpath("string(//title)") or "").strip()
        desc = (tree.xpath("string(//meta[@name='description']/@content)") or "").strip()
        h1 = [x.text_content().strip() for x in tree.xpath("//h1")]
        canonical = tree.xpath("//link[@rel='canonical']/@href")
        schemas = tree.xpath("//script[@type='application/ld+json']")
        words = len(" ".join(tree.xpath("//body//text()[not(ancestor::script) and not(ancestor::style)]")).split())
        rel = path.relative_to(DIST)
        url = "/" if rel.name == "index.html" else "/" + str(rel).replace("/index.html", "").replace(".html", "")
        pages.append({
            "url": url, "title": title, "title_len": len(title), "desc_len": len(desc),
            "h1": len(h1), "canonical": bool(canonical), "schemas": len(schemas), "words": words
        })
    return pages


PAGES = page_audit()

seo_rows = [
    ("Titres", "Titres uniques et descriptifs sur les pages clés. Plusieurs dépassent toutefois la plage idéale de 50-60 caractères.", "Good"),
    ("Meta descriptions", "Présentes sur l'ensemble des pages statiques auditées, avec une promesse et un vocabulaire métier clairs.", "Good"),
    ("H1 et hiérarchie", "Un H1 unique sur les pages publiques principales; hiérarchie H2/H3 globalement cohérente.", "Good"),
    ("Canonical", "Balise canonical générée de façon centralisée par les layouts.", "Good"),
    ("Robots et sitemap", "robots.txt autorise le crawl et exclut /admin; sitemap dynamique présent. Le sitemap public a toutefois été bloqué par le client de navigation lors du contrôle.", "Needs Attention"),
    ("Indexation observée", "La recherche site:elevenlease.fr n'a retourné aucune page Eleven Lease pendant l'audit. À vérifier immédiatement dans Google Search Console.", "Needs Attention"),
    ("Réseaux sociaux", "Open Graph et Twitter Cards présents, mais l'image sociale est le logo carré plutôt qu'un visuel 1200 x 630 conçu pour le clic.", "Needs Attention"),
    ("Images", "Alt text produit correct. Les images véhicules distantes n'avaient pas encore de dimensions naturelles au contrôle, ce qui peut favoriser les décalages de mise en page.", "Needs Attention"),
    ("Contenu éditorial", "Huit articles de 694 à 1 001 mots et un guide LOA complet: base éditoriale solide.", "Good"),
    ("Maillage interne", "Navigation, guide, articles et CTA sont reliés; renforcer les liens contextuels entre articles et fiches véhicules.", "Needs Attention"),
]

geo_rows = [
    ("Entité de marque", "Organization schema, nom, logo, URL, email et profils sociaux sont déclarés de manière cohérente.", "Good"),
    ("À propos", "Page dédiée avec date de création, expérience sectorielle, 400+ dossiers et une dizaine de partenaires.", "Good"),
    ("Preuves vérifiables", "Les chiffres 400+, 24h, 100% d'offres vérifiées et taux d'acceptation supérieur ne sont accompagnés d'aucune méthode, période ou source.", "Needs Attention"),
    ("Identité légale", "Forme juridique, siège, SIREN, RCS, capital, TVA et directeur de publication sont encore indiqués « À compléter ».", "Missing"),
    ("Cohérence hébergement", "Mentions légales: Vercel. Politique de confidentialité: Netlify. Cette contradiction fragilise la confiance.", "Missing"),
    ("Équipe et auteurs", "Aucun nom, portrait, fonction ni profil professionnel vérifiable sur la page À propos ou les articles.", "Needs Attention"),
    ("Partenaires", "Le site évoque une dizaine de partenaires mais n'en nomme aucun et ne montre aucun logo ou cadre de collaboration.", "Needs Attention"),
    ("Citations externes", "Les articles pédagogiques ne citent presque pas de sources publiques ou réglementaires, alors que la LOA est un sujet financier sensible.", "Needs Attention"),
    ("Originalité", "Positionnement distinctif sur le kilométrage illimité et les gros rouleurs, exprimé de façon constante.", "Good"),
]

aeo_rows = [
    ("Réponses directes", "Le guide LOA répond directement aux questions principales avec des sections bien nommées.", "Good"),
    ("FAQ", "FAQ visible et FAQPage schema sur l'accueil, le guide LOA et la page contact.", "Good"),
    ("HowTo", "Parcours en quatre étapes et HowTo schema sur l'accueil.", "Good"),
    ("Formats extractibles", "Listes, étapes et comparaison Eleven Lease / financement classique sont faciles à synthétiser.", "Good"),
    ("Questions longue traîne", "Les articles couvrent jeunes conducteurs, indépendants, apport, électrique, assurance, résiliation et restitution.", "Good"),
    ("Tableaux comparatifs", "Peu de vrais tableaux sémantiques dans les articles; potentiel sur LOA vs LLD, avec/sans apport et coûts.", "Needs Attention"),
    ("Réponses sourcées", "Les réponses sont claires mais manquent de références officielles et de dates de mise à jour visibles.", "Needs Attention"),
]

priorities = [
    ("Critique", "Finaliser les mentions légales, l'identité du responsable de traitement et les informations d'intermédiation applicables.", "GEO / Confiance", "Moyen", "Très élevé"),
    ("Critique", "Corriger la contradiction Vercel / Netlify et préciser durées, destinataires, transferts et consentement pour les données financières.", "Conformité", "Faible à moyen", "Très élevé"),
    ("Haute", "Vérifier l'indexation dans Search Console, soumettre le sitemap et demander l'indexation des pages piliers.", "SEO", "Faible", "Très élevé"),
    ("Haute", "Rendre les preuves commerciales vérifiables: période, définition et méthodologie pour 400+, 24h, 100% et taux d'acceptation.", "GEO / Conversion", "Moyen", "Élevé"),
    ("Haute", "Ajouter des personnes réelles: noms, rôles, expérience, portrait et profils professionnels; signer les articles.", "GEO / E-E-A-T", "Moyen", "Élevé"),
    ("Haute", "Corriger le débordement horizontal mobile observé à 390 px (scrollWidth 429 px).", "UX mobile", "Faible", "Élevé"),
    ("Moyenne", "Créer des visuels Open Graph 1200 x 630 par type de page et par article.", "SEO social", "Moyen", "Moyen"),
    ("Moyenne", "Ajouter des sources officielles datées dans les guides: Service-Public, DGCCRF, Banque de France, CNIL.", "GEO / AEO", "Moyen", "Élevé"),
    ("Moyenne", "Ajouter des tableaux HTML comparatifs et des blocs de réponse de 40-60 mots sous les questions clés.", "AEO", "Moyen", "Moyen"),
    ("Quick Win", "Remplacer l'exemple Peugeot 3008 à 359 EUR s'il n'existe plus dans le catalogue dynamique.", "Exactitude", "Faible", "Moyen"),
    ("Quick Win", "Définir width/height ou aspect-ratio explicites pour les images véhicules afin de limiter le CLS.", "Technique", "Faible", "Moyen"),
    ("Quick Win", "Afficher une date de mise à jour et un auteur sur chaque article.", "GEO / SEO", "Faible", "Moyen"),
]


def configure_styles(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12
    for name, size, color, before, after in [
        ("Heading 1", 24, NAVY, 14, 8),
        ("Heading 2", 17, BLUE, 12, 6),
        ("Heading 3", 13, NAVY, 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Arial"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)


def add_running_header_footer(section):
    section.header.is_linked_to_previous = False
    p = section.header.paragraphs[0]
    p.text = "elevenlease.fr                                      SEO / GEO / AEO Audit Report"
    p.style = "Header"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.name = "Arial"; r.font.size = Pt(8.5); r.font.color.rgb = RGBColor.from_string(NAVY)
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Eleven Lease - Audit du 24 juillet 2026   |   ")
    r.font.name = "Arial"; r.font.size = Pt(8); r.font.color.rgb = RGBColor.from_string("64748B")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    p._p.append(fld)


def add_score_table(doc):
    rows = [
        ("SEO", "6.5/10", "Fondations solides, indexation à contrôler"),
        ("GEO", "5/10", "Marque claire, confiance et preuves insuffisantes"),
        ("AEO", "7.5/10", "FAQ, HowTo et longue traîne déjà bien structurés"),
        ("Combiné", "19/30", "Bon produit éditorial, crédibilité à consolider"),
    ]
    add_table(doc, ["Dimension", "Score", "Lecture"], rows, [1.2, 1.0, 4.3])


doc = Document()
configure_styles(doc)

# Cover
sec = doc.sections[0]
sec.top_margin = sec.bottom_margin = Inches(0.6)
for _ in range(4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(22)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("elevenlease.fr")
r.bold = True; r.font.name = "Arial"; r.font.size = Pt(36); r.font.color.rgb = RGBColor.from_string(NAVY)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("SEO / GEO / AEO Audit Report")
r.font.name = "Arial"; r.font.size = Pt(18); r.font.color.rgb = RGBColor.from_string(BLUE)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("AUDIT COMPLET  |  24 JUILLET 2026")
r.bold = True; r.font.name = "Arial"; r.font.size = Pt(10); r.font.color.rgb = RGBColor.from_string("64748B")
doc.add_paragraph()
cover = doc.add_table(rows=1, cols=3)
cover.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (label, score, color) in enumerate((("SEO", "6.5", AMBER), ("GEO", "5", AMBER), ("AEO", "7.5", AMBER))):
    cell = cover.cell(0, i); shade(cell, color); margins(cell, 220, 120, 220, 120)
    set_cell_text(cell, f"{label}\n{score}/10\nÀ renforcer", bold=True, color=WHITE, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
sec = doc.add_section(WD_SECTION.NEW_PAGE)
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
add_running_header_footer(sec)
doc.sections[0].header.paragraphs[0].text = ""
doc.sections[0].footer.paragraphs[0].text = ""

doc.add_heading("Executive Summary", level=1)
t = doc.add_table(rows=1, cols=1)
shade(t.cell(0, 0), LIGHT_BLUE)
set_cell_text(t.cell(0, 0),
    "Eleven Lease dispose d'une proposition de valeur remarquable - le leasing à kilométrage illimité - et d'une base technique SEO déjà supérieure à celle de nombreux jeunes sites: métadonnées, canonical, données structurées, guide LOA et huit articles spécialisés. Le principal frein n'est pas le design mais la confiance: le site collecte des informations personnelles et financières alors que l'identité légale reste incomplète et que les informations d'hébergement se contredisent. La visibilité organique semble également limitée: aucune page Eleven Lease n'est apparue dans la recherche site:elevenlease.fr effectuée pendant l'audit. La priorité est donc de sécuriser la conformité et les preuves, puis d'accélérer l'indexation et l'autorité éditoriale.",
    size=10)
doc.add_paragraph()
add_score_table(doc)

doc.add_heading("Top 6 des modifications recommandées", level=2)
for text in [
    "Finaliser immédiatement les mentions légales et la politique de confidentialité.",
    "Contrôler Google Search Console, soumettre le sitemap et demander l'indexation des pages piliers.",
    "Documenter toutes les promesses chiffrées et afficher de vraies preuves.",
    "Humaniser la marque avec une équipe identifiable et des auteurs signés.",
    "Corriger le débordement horizontal mobile mesuré sur l'accueil.",
    "Transformer les guides en contenus sourcés et citables par les moteurs d'IA.",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)

doc.add_heading("Pages Audited", level=1)
page_rows = []
for p in PAGES:
    kind = "Article" if p["url"].startswith("/articles/") else "Page publique"
    notes = f"{p['words']} mots; {p['h1']} H1; {p['schemas']} schema; title {p['title_len']} car.; meta {p['desc_len']} car."
    page_rows.append((p["url"], kind, notes))
add_table(doc, ["URL", "Type", "Notes"], page_rows, [2.2, 1.0, 3.3])

doc.add_heading("SEO Analysis", level=1)
doc.add_paragraph("Score: 6.5/10. La base on-page est saine; le principal risque est l'indexation réelle et l'absence de validation Search Console.")
add_table(doc, ["Signal", "Finding", "Status"], seo_rows, [1.35, 4.25, 0.9], status_col=2)

doc.add_heading("GEO Analysis", level=1)
doc.add_paragraph("Score: 5/10. La marque est compréhensible par les moteurs d'IA, mais manque encore de preuves, d'identité humaine et de cohérence réglementaire pour être citée avec confiance.")
add_table(doc, ["Signal", "Finding", "Status"], geo_rows, [1.35, 4.25, 0.9], status_col=2)

doc.add_heading("AEO Analysis", level=1)
doc.add_paragraph("Score: 7.5/10. Le site est déjà bien structuré pour les réponses courtes et les questions longue traîne; le prochain levier est la qualité des sources et les formats comparatifs.")
add_table(doc, ["Signal", "Finding", "Status"], aeo_rows, [1.35, 4.25, 0.9], status_col=2)

doc.add_heading("UX, mobile et conversion", level=1)
ux = [
    ("Proposition de valeur", "Le bénéfice est compris dès le H1 et les deux CTA sont visibles immédiatement.", "Good"),
    ("Identité visuelle", "Univers premium, typographie et accent rose distinctifs; cohérence forte entre les sections.", "Good"),
    ("Mobile", "À 390 px, documentElement.scrollWidth = 429 px: un élément dépasse de 39 px et peut provoquer un défilement horizontal.", "Needs Attention"),
    ("Formulaire", "Parcours détaillé et rassurant, mais la demande d'âge, revenus, charges et statut FICP exige une confiance juridique irréprochable.", "Needs Attention"),
    ("Catalogue", "Cartes claires et loyers lisibles. Les conditions essentielles devraient être plus visibles avant la simulation.", "Needs Attention"),
    ("Preuve sociale", "Témoignages présents, mais sans lien vers une plateforme, date, véhicule ou preuve d'achat.", "Needs Attention"),
    ("Accessibilité", "Labels, aria-labels, focus et réduction de mouvement sont présents dans plusieurs composants.", "Good"),
]
add_table(doc, ["Signal", "Finding", "Status"], ux, [1.35, 4.25, 0.9], status_col=2)

doc.add_heading("Priority Recommendations", level=1)
add_table(doc, ["Priorité", "Issue", "Dimension", "Effort", "Impact"], priorities, [0.8, 3.1, 1.05, 0.8, 0.75])

doc.add_heading("Plan d'action conseillé", level=1)
for title, body in [
    ("Semaine 1 - sécuriser", "Mentions légales, politique de confidentialité, consentement des formulaires, cohérence Vercel/Netlify, audit mobile et Search Console."),
    ("Semaines 2-3 - crédibiliser", "Page équipe, auteurs, méthodologie des chiffres, preuves partenaires et témoignages vérifiables."),
    ("Mois 2 - gagner en visibilité", "Sources officielles, visuels sociaux, tableaux comparatifs, maillage articles-vers-offres et contenus par profil."),
    ("Mois 3 - mesurer", "Suivi Search Console, conversions formulaire/WhatsApp, requêtes par page, taux d'indexation et Core Web Vitals."),
]:
    doc.add_heading(title, level=2)
    doc.add_paragraph(body)

doc.add_heading("What's Working Well", level=1)
strengths = [
    ("Positionnement", "Le kilométrage illimité donne une vraie différence et cible un usage précis: les gros rouleurs."),
    ("Architecture", "Accueil, catalogue, guide LOA, articles, simulation, contact et À propos forment un parcours cohérent."),
    ("Données structurées", "Organization, FAQPage, HowTo, BreadcrumbList, CollectionPage, Article et Vehicle sont déjà exploités."),
    ("Contenu", "Huit articles substantiels et un guide LOA complet donnent une base rare pour un site récent."),
    ("Conversion", "CTA répétés sans être opaques, WhatsApp visible et promesse de réponse sous 24h."),
    ("Technique", "Build Astro statique rapide, canonical centralisé, robots et sitemap prévus dans le projet."),
]
add_table(doc, ["Force", "Evidence"], strengths, [1.4, 5.1])

doc.add_heading("Glossaire", level=1)
doc.add_paragraph("SEO - optimisation pour les moteurs de recherche traditionnels: crawl, indexation, pertinence et expérience.")
doc.add_paragraph("GEO - optimisation pour les moteurs génératifs: clarté de l'entité, preuves, expertise et contenu facile à citer.")
doc.add_paragraph("AEO - optimisation pour les réponses directes: FAQ, extraits enrichis, questions naturelles, listes et tableaux.")
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(3)
p.paragraph_format.space_after = Pt(0)
r = p.add_run("Limites: aucun accès à Search Console, Analytics, aux conversions ou aux backlinks. Confirmer les Core Web Vitals avec PageSpeed Insights et Search Console.")
r.font.name = "Arial"
r.font.size = Pt(8.5)
r.font.color.rgb = RGBColor.from_string("64748B")

doc.core_properties.title = "Audit SEO / GEO / AEO - elevenlease.fr"
doc.core_properties.subject = "Audit complet du site Eleven Lease"
doc.core_properties.author = "Codex"
doc.save(DOCX)
print(DOCX)
